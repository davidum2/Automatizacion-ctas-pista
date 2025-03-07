import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crearXML(template_path, output_dir, data):
    """
    Crea un documento Word con el contenido del XML.
    
    Args:
        template_path (str): Ruta a la plantilla para el documento XML
        output_dir (str): Directorio donde se guardará el documento
        data (dict): Datos de la factura
        
    Returns:
        str: Ruta al documento generado
    """
    try:
        # Verificar que la plantilla existe
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"No se encontró la plantilla: {template_path}")
        
        # Cargar la plantilla
        doc = Document(template_path)
        
        # Añadir título
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_run = titulo.add_run("COMPROBANTE FISCAL DIGITAL (CFDI)")
        titulo_run.bold = True
        titulo_run.font.size = Pt(14)
        
        # Añadir información básica
        info = doc.add_paragraph()
        info_run = info.add_run(f"""
Folio Fiscal (UUID): {data['Folio_Fiscal']}
Serie y Folio: {data['Serie']}{data['Numero']}
Fecha de Emisión: {data['Fecha_factura']}
""")
        info_run.font.size = Pt(11)
        
        # Añadir datos del emisor
        emisor = doc.add_paragraph()
        emisor_run = emisor.add_run("DATOS DEL EMISOR:")
        emisor_run.bold = True
        emisor_run.font.size = Pt(12)
        
        emisor_info = doc.add_paragraph()
        emisor_info_run = emisor_info.add_run(f"""
Nombre: {data['Emisor']['Nombre']}
RFC: {data['Emisor']['Rfc']}
""")
        emisor_info_run.font.size = Pt(11)
        
        # Añadir datos del receptor
        receptor = doc.add_paragraph()
        receptor_run = receptor.add_run("DATOS DEL RECEPTOR:")
        receptor_run.bold = True
        receptor_run.font.size = Pt(12)
        
        receptor_info = doc.add_paragraph()
        receptor_info_run = receptor_info.add_run(f"""
Nombre: {data['Receptor']['Nombre']}
RFC: {data['Receptor']['Rfc']}
""")
        receptor_info_run.font.size = Pt(11)
        
        # Añadir conceptos
        conceptos = doc.add_paragraph()
        conceptos_run = conceptos.add_run("CONCEPTOS:")
        conceptos_run.bold = True
        conceptos_run.font.size = Pt(12)
        
        # Crear tabla para conceptos
        tabla = doc.add_table(rows=1, cols=3)
        tabla.style = 'Table Grid'
        
        # Añadir encabezados
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = "Descripción"
        hdr_cells[1].text = "Cantidad"
        hdr_cells[2].text = "Valor"
        
        # Añadir filas para cada concepto
        for descripcion, cantidad in data['Conceptos'].items():
            row_cells = tabla.add_row().cells
            row_cells[0].text = descripcion
            row_cells[1].text = str(cantidad)
            row_cells[2].text = ""  # No tenemos el valor unitario
        
        # Añadir total
        total = doc.add_paragraph()
        total_run = total.add_run(f"TOTAL: {data['monto']}")
        total_run.bold = True
        total_run.font.size = Pt(12)
        
        # Añadir nota
        nota = doc.add_paragraph()
        nota.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        nota_run = nota.add_run("NOTA: Este documento es una representación impresa del Comprobante Fiscal Digital (CFDI) y tiene validez fiscal.")
        nota_run.italic = True
        nota_run.font.size = Pt(9)
        
        # Guardar el documento
        output_path = os.path.join(output_dir, "xml.docx")
        doc.save(output_path)
        
        return output_path
        
    except Exception as e:
        raise Exception(f"Error al crear documento XML: {str(e)}")
