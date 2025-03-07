import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def legalizacionVerificacion(template_path, output_dir, data):
    """
    Crea un documento de legalización de verificación del SAT basado en una plantilla.
    
    Args:
        template_path (str): Ruta a la plantilla de legalización
        output_dir (str): Directorio donde se guardará el documento
        data (dict): Datos de la factura
        
    Returns:
        str: Ruta al documento generado
    """
    try:
        # Verificar que la plantilla existe
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"No se encontró la plantilla: {template_path}")
        
        # Cargar la plantilla
        doc = Document(template_path)
        
        # Buscar marcadores en el documento y reemplazarlos con los datos
        for paragraph in doc.paragraphs:
            # Reemplazar marcadores en el texto del párrafo
            if '{{FECHA_DOCUMENTO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{FECHA_DOCUMENTO}}', data['Fecha_doc'])
            if '{{SERIE_NUMERO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{SERIE_NUMERO}}', f"{data['Serie']}{data['Numero']}")
            if '{{FECHA_FACTURA}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{FECHA_FACTURA}}', data['Fecha_factura'])
            if '{{NOMBRE_EMISOR}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NOMBRE_EMISOR}}', data['Nombre_Emisor'])
            if '{{MONTO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{MONTO}}', data['monto'])
            if '{{EMPLEO_RECURSO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{EMPLEO_RECURSO}}', data['Empleo_recurso'])
            if '{{NO_PARTIDA}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NO_PARTIDA}}', data['No_partida'])
            if '{{MES}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{MES}}', data['Mes'])
            if '{{NO_MENSAJE}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NO_MENSAJE}}', data['No_mensaje'])
            if '{{FECHA_MENSAJE}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{FECHA_MENSAJE}}', data['Fecha_mensaje'])
            if '{{FOLIO_FISCAL}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{FOLIO_FISCAL}}', data['Folio_Fiscal'])
        
        # También buscar en las tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{{FECHA_DOCUMENTO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{FECHA_DOCUMENTO}}', data['Fecha_doc'])
                        if '{{SERIE_NUMERO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{SERIE_NUMERO}}', f"{data['Serie']}{data['Numero']}")
                        if '{{FECHA_FACTURA}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{FECHA_FACTURA}}', data['Fecha_factura'])
                        if '{{NOMBRE_EMISOR}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{NOMBRE_EMISOR}}', data['Nombre_Emisor'])
                        if '{{MONTO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{MONTO}}', data['monto'])
                        if '{{EMPLEO_RECURSO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{EMPLEO_RECURSO}}', data['Empleo_recurso'])
                        if '{{NO_PARTIDA}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{NO_PARTIDA}}', data['No_partida'])
                        if '{{MES}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{MES}}', data['Mes'])
                        if '{{NO_MENSAJE}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{NO_MENSAJE}}', data['No_mensaje'])
                        if '{{FECHA_MENSAJE}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{FECHA_MENSAJE}}', data['Fecha_mensaje'])
                        if '{{FOLIO_FISCAL}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{FOLIO_FISCAL}}', data['Folio_Fiscal'])
        
        # Guardar el documento
        output_path = os.path.join(output_dir, "legalizacion_verificacion.docx")
        doc.save(output_path)
        
        return output_path
        
    except Exception as e:
        raise Exception(f"Error al crear legalización de verificación: {str(e)}")
