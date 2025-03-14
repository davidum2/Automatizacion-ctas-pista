import os
import re
from decimal import Decimal, InvalidOperation
import logging
from docx import Document
from datetime import datetime
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

# Configurar logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def calcular_montos_facturas(facturas_info):
    """
    Calcula información resumida de las facturas procesadas.

    Args:
        facturas_info (list): Lista de facturas procesadas

    Returns:
        dict: Diccionario con información resumida
    """
    # Inicializar contadores
    total_facturas = 0
    monto_total = Decimal('0.00')
    montos_individuales = []

    # Procesar cada factura
    for factura in facturas_info:
        # Saltear elementos que no son diccionarios
        if not isinstance(factura, dict):
            continue

        total_facturas += 1

        # Primero intentar usar el valor decimal directo
        if 'monto_decimal' in factura:
            monto = factura['monto_decimal']
        else:
            # Extraer el monto y convertirlo a Decimal
            monto_str = factura.get('monto', '0')
            if isinstance(monto_str, str):
                # Limpiar el string de monto (eliminar símbolos de moneda y separadores de miles)
                monto_limpio = re.sub(r'[^\d.]', '', monto_str.replace(',', ''))
                try:
                    monto = Decimal(monto_limpio)
                except InvalidOperation:
                    logger.warning(f"No se pudo convertir el monto '{monto_str}' a decimal")
                    monto = Decimal('0.00')
            elif isinstance(monto_str, (int, float)):
                monto = Decimal(str(monto_str))
            else:
                monto = Decimal('0.00')
        
        montos_individuales.append(monto)
        monto_total += monto

    # Formatear el monto total como string con formato moneda
    monto_formateado = f"$ {monto_total:,.2f}"

    return {
        'total_facturas': total_facturas,
        'monto_total': monto_total,
        'monto_formateado': monto_formateado,
        'montos_individuales': montos_individuales
    }

def aplicar_formato_geomanist(paragraph):
    """
    Aplica el formato Geomanist tamaño 10 a un párrafo.

    Args:
        paragraph: Párrafo al que aplicar el formato
    """
    for run in paragraph.runs:
        run.font.name = "Geomanist"
        run.font.size = Pt(10)

def aplicar_formato_a_documento(doc):
    """
    Aplica el formato Geomanist tamaño 10 a todo el documento.

    Args:
        doc: Documento Word a formatear
    """
    # Aplicar a párrafos
    for paragraph in doc.paragraphs:
        aplicar_formato_geomanist(paragraph)

    # Aplicar a tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    aplicar_formato_geomanist(paragraph)

def aplicar_bordes_celda(celda):
    """
    Aplica bordes a todos los lados de una celda.

    Args:
        celda: Celda de tabla a la que aplicar bordes
    """
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Crear elemento de bordes
    tcBorders = OxmlElement('w:tcBorders')
    tcPr.append(tcBorders)
    
    # Bordes: superior, inferior, izquierdo, derecho
    for border_type in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)

def aplicar_formato_celda(celda, centrar=True, aplicar_bordes=True, fuente_geomanist=True):
    """
    Aplica formato completo a una celda: bordes, alineación y fuente.

    Args:
        celda: Celda de tabla
        centrar: Si el texto debe centrarse horizontalmente
        aplicar_bordes: Si se deben aplicar bordes
        fuente_geomanist: Si se debe aplicar fuente Geomanist
    """
    # Aplicar bordes
    if aplicar_bordes:
        aplicar_bordes_celda(celda)
    
    # Aplicar alineación vertical al centro
    celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Centrar texto horizontalmente en todos los párrafos de la celda
    for paragraph in celda.paragraphs:
        if centrar:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Aplicar formato Geomanist
        if fuente_geomanist:
            for run in paragraph.runs:
                run.font.name = "Geomanist"
                run.font.size = Pt(10)

def reemplazar_marcadores_texto(texto, reemplazos):
    """
    Reemplaza marcadores en un texto.

    Args:
        texto: Texto con marcadores
        reemplazos: Diccionario con los marcadores y sus reemplazos

    Returns:
        str: Texto con los marcadores reemplazados
    """
    texto_resultado = texto
    for key, value in reemplazos.items():
        if key in texto_resultado:
            texto_resultado = texto_resultado.replace(key, str(value))
    return texto_resultado

def reemplazar_marcadores_en_run(run, reemplazos):
    """
    Reemplaza marcadores en un run individual.

    Args:
        run: Run que puede contener marcadores
        reemplazos: Diccionario con los marcadores y sus reemplazos

    Returns:
        bool: True si se realizó algún reemplazo, False en caso contrario
    """
    texto_original = run.text
    texto_nuevo = reemplazar_marcadores_texto(texto_original, reemplazos)
    
    if texto_nuevo != texto_original:
        run.text = texto_nuevo
        return True
    
    return False

def reemplazar_marcadores_en_parrafo(paragraph, reemplazos):
    """
    Reemplaza marcadores en un párrafo completo, manteniendo el formato de cada run.

    Args:
        paragraph: Párrafo que puede contener marcadores
        reemplazos: Diccionario con los marcadores y sus reemplazos

    Returns:
        bool: True si se realizó algún reemplazo, False en caso contrario
    """
    texto_completo = paragraph.text
    
    # Verificar si hay marcadores en el párrafo completo
    tiene_marcadores = any(key in texto_completo for key in reemplazos)
    
    if not tiene_marcadores:
        return False
    
    # Si hay marcadores, reemplazar en cada run
    se_realizo_reemplazo = False
    for run in paragraph.runs:
        if reemplazar_marcadores_en_run(run, reemplazos):
            se_realizo_reemplazo = True
    
    # Si el reemplazo run por run no funcionó, probar con método alternativo
    if not se_realizo_reemplazo and tiene_marcadores:
        # Guardar el formato de cada run
        runs_info = []
        for run in paragraph.runs:
            runs_info.append({
                'texto': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'color': run.font.color.rgb if run.font.color else None
            })
        
        # Reemplazar en el texto completo
        texto_nuevo = reemplazar_marcadores_texto(texto_completo, reemplazos)
        
        # Si hay cambios, aplicar el texto nuevo
        if texto_nuevo != texto_completo:
            # Limpiar el párrafo
            for _ in range(len(paragraph.runs)):
                paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
            
            # Crear un nuevo run con el texto completo y recuperar formato del primer run original
            if runs_info:
                run = paragraph.add_run(texto_nuevo)
                run.bold = runs_info[0]['bold']
                run.italic = runs_info[0]['italic']
                run.underline = runs_info[0]['underline']
                if runs_info[0]['font_name']:
                    run.font.name = runs_info[0]['font_name']
                if runs_info[0]['font_size']:
                    run.font.size = runs_info[0]['font_size']
                if runs_info[0]['color']:
                    run.font.color.rgb = runs_info[0]['color']
            else:
                paragraph.add_run(texto_nuevo)
            
            se_realizo_reemplazo = True
    
    return se_realizo_reemplazo

def reemplazar_marcadores_en_celda(cell, reemplazos):
    """
    Reemplaza marcadores en todos los párrafos de una celda.

    Args:
        cell: Celda de tabla que puede contener marcadores
        reemplazos: Diccionario con los marcadores y sus reemplazos

    Returns:
        bool: True si se realizó algún reemplazo, False en caso contrario
    """
    se_realizo_reemplazo = False
    
    for paragraph in cell.paragraphs:
        if reemplazar_marcadores_en_parrafo(paragraph, reemplazos):
            se_realizo_reemplazo = True
    
    return se_realizo_reemplazo

def reemplazar_marcadores_en_documento(doc, reemplazos):
    """
    Reemplaza todos los marcadores en un documento completo.
    Esta función maneja tanto párrafos como tablas y preserva el formato.

    Args:
        doc: Documento Word
        reemplazos: Diccionario con los marcadores y sus reemplazos
    """
    # Procesar todos los párrafos del documento
    for paragraph in doc.paragraphs:
        reemplazar_marcadores_en_parrafo(paragraph, reemplazos)
    
    # Procesar todas las tablas del documento
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_marcadores_en_celda(cell, reemplazos)

def encontrar_plantilla(nombre_archivo, base_dir=None):
    """
    Busca una plantilla en diferentes ubicaciones posibles.

    Args:
        nombre_archivo: Nombre del archivo de plantilla
        base_dir: Directorio base opcional

    Returns:
        str: Ruta completa a la plantilla o None si no se encuentra
    """
    if not base_dir:
        base_dir = os.path.dirname(os.path.abspath(__file__))

    # Lista de posibles ubicaciones de plantillas
    directorios_posibles = [
        os.path.join(os.path.dirname(base_dir), "plantillas"),
        os.path.join(os.path.dirname(base_dir), "templates"),
        os.path.join(base_dir, "plantillas"),
        os.path.join(base_dir, "templates")
    ]

    # Buscar el archivo en cada directorio
    for directorio in directorios_posibles:
        ruta_posible = os.path.join(directorio, nombre_archivo)
        if os.path.exists(ruta_posible):
            return ruta_posible

    return None

def procesar_plantillas_partida(partida, facturas_info, partida_dir, datos_comunes):
    """
    Procesa todas las plantillas para una partida específica.
    """
    logger.info(f"Procesando plantillas para partida {partida.get('numero', 'desconocida')}")

    # Archivos generados
    archivos_generados = {}

    try:
        # Si la información de facturas ya está calculada, usarla
        if 'info_facturas' in datos_comunes:
            info_facturas = datos_comunes['info_facturas']
            logger.info(f"Usando información de facturas proporcionada - Total: {info_facturas['monto_formateado']}")
        else:
            # Calcular información resumida de facturas
            info_facturas = calcular_montos_facturas(facturas_info)
            logger.info(f"Calculados totales para {info_facturas['total_facturas']} facturas. "
                       f"Monto total: {info_facturas['monto_formateado']}")
            # Añadir la información resumida a los datos comunes
            datos_comunes['info_facturas'] = info_facturas

        # Procesar plantilla de ingresos-egresos
        ruta_ingresos = procesar_plantilla_ingresos(
            partida_dir,
            partida,
            facturas_info,
            datos_comunes
        )
        archivos_generados["ingresos"] = ruta_ingresos
        logger.info(f"Plantilla de ingresos generada en: {ruta_ingresos}")

        # Procesar plantilla de relación de facturas
        ruta_facturas = procesar_plantilla_facturas(
            partida_dir,
            partida,
            facturas_info,
            datos_comunes
        )
        archivos_generados["facturas"] = ruta_facturas
        logger.info(f"Plantilla de facturas generada en: {ruta_facturas}")

        # Procesar plantilla de oficio
        ruta_oficio = procesar_plantilla_oficio(
            partida_dir,
            partida,
            facturas_info,
            datos_comunes
        )
        archivos_generados["oficio"] = ruta_oficio
        logger.info(f"Plantilla de oficio generada en: {ruta_oficio}")

        return archivos_generados

    except Exception as e:
        logger.error(f"Error al procesar plantillas de partida: {str(e)}")
        raise Exception(f"Error al procesar plantillas de partida: {str(e)}")

# aqui van las funciones de cada plantilla

def procesar_plantilla_ingresos(output_dir, partida, facturas_info, datos_comunes):
    """
    Procesa la plantilla de ingresos/egresos en formato Word.
    """
    try:
        # Buscar la plantilla
        template_path = encontrar_plantilla("ingresos_egresos.docx")
        if not template_path:
            raise FileNotFoundError("No se encontró la plantilla de ingresos/egresos")

        logger.info(f"Utilizando plantilla: {template_path}")

        # Cargar la plantilla
        doc = Document(template_path)

        # Datos comunes
        mes = datos_comunes.get('mes_asignado', '').capitalize()
        partida_num = partida.get('numero', '')
        descripcion = partida.get('descripcion', '')
        fecha_doc = datos_comunes.get('fecha_documento_texto', '')

        # Información de facturas
        info_facturas = datos_comunes.get('info_facturas', {})
        total_facturas = info_facturas.get('total_facturas', 0)
        monto_total = info_facturas.get('monto_total', Decimal('0.00'))
        monto_formateado = info_facturas.get('monto_formateado', "$ 0.00")

        # Calcular montos específicos
        monto_asignado = Decimal(str(partida.get('monto', 0)))
        monto_asignado_str = f"$ {monto_asignado:,.2f}"
        aportacion = monto_total - monto_asignado
        aportacion_str = f"$ {aportacion:,.2f}"
        suma_ingresos = monto_asignado + aportacion
        suma_ingresos_str = f"$ {suma_ingresos:,.2f}"
        saldo = suma_ingresos - monto_total
        saldo_str = f"$ {saldo:,.2f}"

        # Datos del personal
        personal_vobo = datos_comunes.get('personal_vobo', {})
        grado_vobo = personal_vobo.get('Grado_Vo_Bo', '')
        nombre_vobo = personal_vobo.get('Nombre_Vo_Bo', '')
        matricula_vobo = personal_vobo.get('Matricula_Vo_Bo', '')

        # Crear diccionario de reemplazos
        reemplazos = {
            '{{FECHA_DOCUMENTO}}': fecha_doc,
            '{{MES}}': mes,
            '{{PARTIDA}}': partida_num,
            '{{DESCRIPCION}}': descripcion,
            '{{TOTAL_FACTURAS}}': str(total_facturas),
            '{{MONTO_TOTAL}}': monto_formateado,
            '{{MONTO}}': monto_asignado_str,
            '{{APORTACION}}': aportacion_str,
            '{{SUMA_INGRESOS}}': suma_ingresos_str,
            '{{EGRESOS}}': monto_formateado,
            '{{SALDO}}': saldo_str,
            '{{GRADO_VO_BO}}': grado_vobo,
            '{{NOMBRE_VO_BO}}': nombre_vobo,
            '{{MATRICULA_VO_BO}}': matricula_vobo
        }

        # Reemplazar todos los marcadores utilizando la función mejorada
        reemplazar_marcadores_en_documento(doc, reemplazos)

        # Realizar una verificación adicional para el último párrafo (firma)
        # Esta verificación asegura que los marcadores en ese párrafo específico sean reemplazados
        if len(doc.paragraphs) > 0:
            # Verificar los últimos párrafos (para asegurar que los de firma sean procesados)
            for i in range(min(5, len(doc.paragraphs))):
                idx = len(doc.paragraphs) - 1 - i
                parrafo = doc.paragraphs[idx]
                
                # Buscar marcadores específicos de firma en este párrafo
                texto_original = parrafo.text
                tiene_marcador = any(key in texto_original for key in ['{{GRADO_VO_BO}}', '{{NOMBRE_VO_BO}}', '{{MATRICULA_VO_BO}}'])
                
                if tiene_marcador:
                    logger.info(f"Encontrado párrafo de firma en posición {idx}")
                    
                    # Reemplazar directamente en el texto y preservar el formato
                    texto_nuevo = texto_original
                    for key, value in reemplazos.items():
                        if key in texto_nuevo:
                            texto_nuevo = texto_nuevo.replace(key, str(value))
                    
                    if texto_nuevo != texto_original:
                        # Guardar el formato (del primer run como referencia)
                        formato = None
                        if parrafo.runs:
                            run_ref = parrafo.runs[0]
                            formato = {
                                'bold': run_ref.bold,
                                'italic': run_ref.italic,
                                'underline': run_ref.underline,
                                'font_name': run_ref.font.name,
                                'font_size': run_ref.font.size
                            }
                        
                        # Limpiar el párrafo
                        for _ in range(len(parrafo.runs)):
                            parrafo.runs[0]._element.getparent().remove(parrafo.runs[0]._element)
                        
                        # Agregar nuevo run con el texto actualizado
                        run_nuevo = parrafo.add_run(texto_nuevo)
                        
                        # Aplicar formato guardado
                        if formato:
                            run_nuevo.bold = formato['bold']
                            run_nuevo.italic = formato['italic']
                            run_nuevo.underline = formato['underline']
                            if formato['font_name']:
                                run_nuevo.font.name = formato['font_name']
                            if formato['font_size']:
                                run_nuevo.font.size = formato['font_size']

        # Realizar una verificación final para asegurarse que los marcadores se hayan reemplazado
        for paragraph in doc.paragraphs:
            for key in reemplazos.keys():
                if key in paragraph.text:
                    logger.warning(f"Marcador {key} no reemplazado en plantilla de ingresos. Intentando nuevo reemplazo.")
                    paragraph.text = paragraph.text.replace(key, str(reemplazos[key]))

        # Aplicar formato Geomanist 10pt a todo el documento
        aplicar_formato_a_documento(doc)

        # Guardar el documento
        output_path = os.path.join(output_dir, f"Ingresos_Egresos_Partida_{partida_num}.docx")
        doc.save(output_path)

        logger.info(f"Documento de ingresos/egresos generado: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"Error al procesar plantilla de ingresos: {str(e)}")
        raise Exception(f"Error al procesar plantilla de ingresos: {str(e)}")




def procesar_plantilla_facturas(output_dir, partida, facturas_info, datos_comunes):
    """
    Procesa la plantilla de relación de facturas en formato Word.
    """
    try:
        # Buscar la plantilla
        template_path = encontrar_plantilla("relcion_facturas.docx")
        if not template_path:
            raise FileNotFoundError("No se encontró la plantilla de facturas")

        logger.info(f"Utilizando plantilla: {template_path}")

        # Cargar la plantilla
        doc = Document(template_path)

        # Datos comunes
        mes = datos_comunes.get('mes_asignado', '').capitalize()
        partida_num = partida.get('numero', '')
        descripcion = partida.get('descripcion', '')
        fecha_doc = datos_comunes.get('fecha_documento_texto', '')

        # Información de facturas
        info_facturas = datos_comunes.get('info_facturas', {})
        total_facturas = info_facturas.get('total_facturas', 0)
        monto_total = info_facturas.get('monto_total', Decimal('0.00'))
        monto_formateado = info_facturas.get('monto_formateado', "$ 0.00")

        # Datos del personal
        personal_vobo = datos_comunes.get('personal_vobo', {})
        grado_vobo = personal_vobo.get('Grado_Vo_Bo', '')
        nombre_vobo = personal_vobo.get('Nombre_Vo_Bo', '')
        matricula_vobo = personal_vobo.get('Matricula_Vo_Bo', '')

        # Reemplazar marcadores de texto en todo el documento
        reemplazos = {
            '{{FECHA_DOCUMENTO}}': fecha_doc,
            '{{MES}}': mes,
            '{{PARTIDA}}': partida_num,
            '{{DESCRIPCION}}': descripcion,
            '{{TOTAL_FACTURAS}}': str(total_facturas),
            '{{MONTO_TOTAL}}': monto_formateado,
            '{{GRADO_VO_BO}}': grado_vobo,
            '{{NOMBRE_VO_BO}}': nombre_vobo,
            '{{MATRICULA_VO_BO}}': matricula_vobo
        }

        # Reemplazar todos los marcadores utilizando la función mejorada
        reemplazar_marcadores_en_documento(doc, reemplazos)

        # Verificar que hay al menos dos tablas en el documento
        if len(doc.tables) < 2:
            raise ValueError("El documento no contiene al menos dos tablas")

        tabla_facturas = doc.tables[1]
        logger.info(f"Se utilizará la segunda tabla con {len(tabla_facturas.rows)} filas y {len(tabla_facturas.columns)} columnas")

        # Filtrar facturas válidas
        facturas_validas = [f for f in facturas_info if isinstance(f, dict)]

        # Si hay más de una fila (encabezado + datos), eliminar todas excepto el encabezado
        while len(tabla_facturas.rows) > 1:
            tr = tabla_facturas._tbl.tr_lst.pop()
            tabla_facturas._tbl.remove(tr)

        # Asegurarse que la fila de encabezado tenga formato adecuado
        if len(tabla_facturas.rows) > 0:
            fila_encabezado = tabla_facturas.rows[0]
            for celda in fila_encabezado.cells:
                aplicar_formato_celda(celda, centrar=True, aplicar_bordes=True, fuente_geomanist=True)

                # Opcional: hacer el texto del encabezado en negrita
                for paragraph in celda.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        # Ahora añadir una fila para cada factura
        for i, factura in enumerate(facturas_validas):
            # Añadir una nueva fila a la tabla
            nueva_fila = tabla_facturas.add_row()
            celdas = nueva_fila.cells

            # Verificar que la tabla tiene las columnas esperadas
            if len(celdas) >= 4:  # Fecha, Número, Emisor, Importe
                # Formatear fecha
                fecha_factura = factura.get('fecha_factura', '')
                if isinstance(fecha_factura, str) and '-' in fecha_factura:
                    try:
                        fecha_obj = datetime.strptime(fecha_factura, '%Y-%m-%d')
                        fecha_factura = fecha_obj.strftime('%d/%m/%Y')
                    except:
                        pass

                # Asignar valores a las celdas
                celdas[0].text = str(fecha_factura)
                celdas[1].text = str(factura.get('serie_numero', ''))
                celdas[2].text = str(factura.get('emisor', ''))
                
                # Usar el valor formateado si existe, o formatear el valor decimal
                if 'monto' in factura and '$' in str(factura['monto']):
                    celdas[3].text = str(factura['monto'])
                elif 'monto_decimal' in factura:
                    celdas[3].text = f"$ {factura['monto_decimal']:,.2f}"
                else:
                    # Intentar formatear lo que haya
                    celdas[3].text = str(factura.get('monto', '0'))

                # Aplicar formato a todas las celdas de la nueva fila
                for celda in celdas:
                    aplicar_formato_celda(celda, centrar=True, aplicar_bordes=True, fuente_geomanist=True)
            else:
                logger.warning(f"La tabla tiene {len(celdas)} columnas, menos de las 4 esperadas")

        # Añadir fila de total si hay facturas
        if facturas_validas:
            fila_total = tabla_facturas.add_row()
            celdas = fila_total.cells

            # Configurar celdas de total
            celdas[0].text = ""
            celdas[1].text = ""
            celdas[2].text = "TOTAL"
            celdas[3].text = monto_formateado

            # Aplicar formato a la fila de total
            # Para las celdas vacías
            aplicar_formato_celda(celdas[0], centrar=True, aplicar_bordes=True, fuente_geomanist=True)
            aplicar_formato_celda(celdas[1], centrar=True, aplicar_bordes=True, fuente_geomanist=True)

            # Para las celdas con "TOTAL" y el monto
            aplicar_formato_celda(celdas[2], centrar=False, aplicar_bordes=True, fuente_geomanist=True)
            celdas[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in celdas[2].paragraphs[0].runs:
                run.bold = True

            aplicar_formato_celda(celdas[3], centrar=False, aplicar_bordes=True, fuente_geomanist=True)
            celdas[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in celdas[3].paragraphs[0].runs:
                run.bold = True

        # Realizar una verificación final para asegurarse que los marcadores se hayan reemplazado
        for paragraph in doc.paragraphs:
            for key in reemplazos.keys():
                if key in paragraph.text:
                    logger.warning(f"Marcador {key} no reemplazado en plantilla de facturas. Intentando nuevo reemplazo.")
                    paragraph.text = paragraph.text.replace(key, str(reemplazos[key]))

        # Guardar el documento
        output_path = os.path.join(output_dir, f"Relacion_Facturas_Partida_{partida.get('numero', '')}.docx")
        doc.save(output_path)

        logger.info(f"Documento de relación de facturas generado: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"Error al procesar plantilla de facturas: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"Error al procesar plantilla de facturas: {str(e)}")
    

def procesar_plantilla_oficio(output_dir, partida, facturas_info, datos_comunes):
    """
    Procesa la plantilla de oficio en formato Word preservando formatos originales.
    Mantiene el tamaño de texto específico de algunos elementos e imágenes.
    """
    try:
        # Buscar la plantilla
        template_path = encontrar_plantilla("Oficio.docx")
        if not template_path:
            raise FileNotFoundError("No se encontró la plantilla de oficio")

        logger.info(f"Utilizando plantilla: {template_path}")

        # Cargar la plantilla
        doc = Document(template_path)

        # Datos comunes
        mes = datos_comunes.get('mes_asignado', '').capitalize()
        partida_num = partida.get('numero', '')
        descripcion = partida.get('descripcion', '')
        fecha_doc = datos_comunes.get('fecha_documento_texto', '')

        # Información de facturas
        info_facturas = datos_comunes.get('info_facturas', {})
        total_facturas = info_facturas.get('total_facturas', 0)
        monto_formateado = info_facturas.get('monto_formateado', "$ 0.00")

        # Datos adicionales para el oficio
        no_oficio = partida.get('numero_adicional', '')

        # Datos del personal
        personal_vobo = datos_comunes.get('personal_vobo', {})
        grado_vobo = personal_vobo.get('Grado_Vo_Bo', '')
        nombre_vobo = personal_vobo.get('Nombre_Vo_Bo', '')
        matricula_vobo = personal_vobo.get('Matricula_Vo_Bo', '')

        # Crear diccionario de reemplazos
        reemplazos = {
            '{{FECHA_DOCUMENTO}}': fecha_doc,
            '{{MES}}': mes,
            '{{PARTIDA}}': partida_num,
            '{{DESCRIPCION}}': descripcion,
            '{{TOTAL_FACTURAS}}': str(total_facturas),
            '{{MONTO_TOTAL}}': monto_formateado,
            '{{NO_OFICIO}}': no_oficio,
            '{{GRADO_VO_BO}}': grado_vobo,
            '{{NOMBRE_VO_BO}}': nombre_vobo,
            '{{MATRICULA_VO_BO}}': matricula_vobo
        }

        # Reemplazar todos los marcadores utilizando la función mejorada
        reemplazar_marcadores_en_documento(doc, reemplazos)

        # Realizar una verificación adicional para el último párrafo (firma)
        # Esta verificación asegura que los marcadores en ese párrafo específico sean reemplazados
        if len(doc.paragraphs) > 0:
            # Verificar los últimos párrafos (para asegurar que los de firma sean procesados)
            for i in range(min(5, len(doc.paragraphs))):
                idx = len(doc.paragraphs) - 1 - i
                parrafo = doc.paragraphs[idx]
                
                # Buscar marcadores específicos de firma en este párrafo
                texto_original = parrafo.text
                tiene_marcador = any(key in texto_original for key in ['{{GRADO_VO_BO}}', '{{NOMBRE_VO_BO}}', '{{MATRICULA_VO_BO}}'])
                
                if tiene_marcador:
                    logger.info(f"Encontrado párrafo de firma en posición {idx}")
                    
                    # Reemplazar directamente en el texto y preservar el formato
                    texto_nuevo = texto_original
                    for key, value in reemplazos.items():
                        if key in texto_nuevo:
                            texto_nuevo = texto_nuevo.replace(key, str(value))
                    
                    if texto_nuevo != texto_original:
                        # Guardar el formato (del primer run como referencia)
                        formato = None
                        if parrafo.runs:
                            run_ref = parrafo.runs[0]
                            formato = {
                                'bold': run_ref.bold,
                                'italic': run_ref.italic,
                                'underline': run_ref.underline,
                                'font_name': run_ref.font.name,
                                'font_size': run_ref.font.size
                            }
                        
                        # Limpiar el párrafo
                        for _ in range(len(parrafo.runs)):
                            parrafo.runs[0]._element.getparent().remove(parrafo.runs[0]._element)
                        
                        # Agregar nuevo run con el texto actualizado
                        run_nuevo = parrafo.add_run(texto_nuevo)
                        
                        # Aplicar formato guardado
                        if formato:
                            run_nuevo.bold = formato['bold']
                            run_nuevo.italic = formato['italic']
                            run_nuevo.underline = formato['underline']
                            if formato['font_name']:
                                run_nuevo.font.name = formato['font_name']
                            if formato['font_size']:
                                run_nuevo.font.size = formato['font_size']

        # Realizar una verificación final para asegurarse que los marcadores se hayan reemplazado
        for paragraph in doc.paragraphs:
            for key in reemplazos.keys():
                if key in paragraph.text:
                    logger.warning(f"Marcador {key} no reemplazado en plantilla de oficio. Intentando nuevo reemplazo.")
                    paragraph.text = paragraph.text.replace(key, str(reemplazos[key]))

        # Guardar el documento
        output_path = os.path.join(output_dir, f"Oficio_Resumen_Partida_{partida_num}.docx")
        doc.save(output_path)

        logger.info(f"Documento de oficio generado: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"Error al procesar plantilla de oficio: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"Error al procesar plantilla de oficio: {str(e)}")