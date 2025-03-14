import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def creacionDocumentos(template_path, output_dir, data, template_name):
    """
    Crea un documento basado en una plantilla.

    Args:
        template_path (str): Ruta a la plantilla
        output_dir (str): Directorio donde se guardará el documento
        data (dict): Datos para rellenar la plantilla
        template_name (str): Nombre de la plantilla (para nombrar el archivo)

    Returns:
        str: Ruta al documento generado
    """
    try:
        # Verificar que la plantilla existe
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"No se encontró la plantilla: {template_path}")

        # Determinar si estamos trabajando con la plantilla XML
        es_plantilla_xml = "xml.docx" in template_path.lower() or template_name.lower() == "xml"

        # Cargar la plantilla
        doc = Document(template_path)

        # Función para aplicar el formato de texto según la plantilla
        def aplicar_formato_texto(paragraph):
            for run in paragraph.runs:
                run.font.name = "Geomanist"
                if es_plantilla_xml:
                    run.font.size = Pt(6)  # Tamaño 6 puntos para la plantilla XML
                else:
                    run.font.size = Pt(10)  # Tamaño predeterminado

        # Buscar marcadores en el documento y reemplazarlos con los datos
        for paragraph in doc.paragraphs:
            # Reemplazar marcadores en el texto del párrafo
            if '{{XML}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{XML}}', data['xml'])
            if '{{FECHA_DOCUMENTO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{FECHA_DOCUMENTO}}', data['Fecha_doc'])
            if '{{SERIE_NUMERO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{SERIE_NUMERO}}', f"{data['Serie']}{data['Numero']}")
            if '{{FECHA_FACTURA}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{FECHA_FACTURA}}', data['Fecha_factura_texto'])
            if '{{PARTIDA}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{PARTIDA}}', data['No_partida'])
            if '{{DESCRIPCION}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{DESCRIPCION}}', data['Descripcion_partida'])
            if '{{NOMBRE_EMISOR}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NOMBRE_EMISOR}}', data['Nombre_Emisor'])
            if '{{MONTO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{MONTO}}', data['monto'])
            if '{{EMPLEO_RECURSO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{EMPLEO_RECURSO}}', data['Empleo_recurso'])
            if '{{MES}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{MES}}', data['Mes'])
            if '{{NO_MENSAJE}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NO_MENSAJE}}', data['No_mensaje'])
            if '{{FECHA_MENSAJE}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{FECHA_MENSAJE}}', data['Fecha_mensaje'])
            if '{{GRADO_RECIBIO_LA_COMPRA}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{GRADO_RECIBIO_LA_COMPRA}}', data['Grado_recibio_la_compra'])
            if '{{NOMBRE_RECIBIO_LA_COMPRA}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NOMBRE_RECIBIO_LA_COMPRA}}', data['Nombre_recibio_la_compra'])
            if '{{MATRICULA_RECIBIO_LA_COMPRA}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{MATRICULA_RECIBIO_LA_COMPRA}}', data['Matricula_recibio_la_compra'])
            if '{{GRADO_VO_BO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{GRADO_VO_BO}}', data['Grado_Vo_Bo'])
            if '{{NOMBRE_VO_BO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{NOMBRE_VO_BO}}', data['Nombre_Vo_Bo'])
            if '{{MATRICULA_VO_BO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{MATRICULA_VO_BO}}', data['Matricula_Vo_Bo'])
            if '{{FOLIO_FISCAL}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{FOLIO_FISCAL}}', data.get('Folio_Fiscal', ''))
            if '{{RFC_EMISOR}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{RFC_EMISOR}}', data.get('Rfc_emisor', ''))
            if '{{RFC_RECEPTOR}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{RFC_RECEPTOR}}', data.get('Rfc_receptor', ''))

            # Aplicar formato después de reemplazar el texto
            aplicar_formato_texto(paragraph)

        # También buscar en las tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{{FECHA_DOCUMENTO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{FECHA_DOCUMENTO}}', data['Fecha_doc'])
                        if '{{SERIE_NUMERO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{SERIE_NUMERO}}', f"{data['Serie']}{data['Numero']}")
                        if '{{FECHA_FACTURA}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{FECHA_FACTURA}}', data['Fecha_factura_texto'])
                        if '{{NOMBRE_EMISOR}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{NOMBRE_EMISOR}}', data['Nombre_Emisor'])
                        if '{{MONTO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{MONTO}}', data['monto'])
                        if '{{PARTIDA}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{PARTIDA}}', data['No_partida'])
                        if '{{DESCRIPCION}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{DESCRIPCION}}', data['Descripcion_partida'])
                        if '{{EMPLEO_RECURSO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{EMPLEO_RECURSO}}', data['Empleo_recurso'])
                        if '{{GRADO_VO_BO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{GRADO_VO_BO}}', data['Grado_Vo_Bo'])
                        if '{{NOMBRE_VO_BO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{NOMBRE_VO_BO}}', data['Nombre_Vo_Bo'])
                        if '{{MATRICULA_VO_BO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{MATRICULA_VO_BO}}', data['Matricula_Vo_Bo'])
                        if '{{MES}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{MES}}', data['Mes'])
                        if '{{NO_MENSAJE}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{NO_MENSAJE}}', data['No_mensaje'])
                        if '{{FECHA_MENSAJE}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{FECHA_MENSAJE}}', data['Fecha_mensaje'])
                        if '{{GRADO_RECIBIO_LA_COMPRA}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{GRADO_RECIBIO_LA_COMPRA}}', data['Grado_recibio_la_compra'])
                        if '{{NOMBRE_RECIBIO_LA_COMPRA}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{NOMBRE_RECIBIO_LA_COMPRA}}', data['Nombre_recibio_la_compra'])
                        if '{{MATRICULA_RECIBIO_LA_COMPRA}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{MATRICULA_RECIBIO_LA_COMPRA}}', data['Matricula_recibio_la_compra'])
                        if '{{FOLIO_FISCAL}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{FOLIO_FISCAL}}', data.get('Folio_Fiscal', ''))
                        if '{{RFC_EMISOR}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{RFC_EMISOR}}', data.get('Rfc_emisor', ''))
                        if '{{RFC_RECEPTOR}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{RFC_RECEPTOR}}', data.get('Rfc_receptor', ''))

                        # Aplicar formato después de reemplazar el texto
                        aplicar_formato_texto(paragraph)

        # Realizar verificación adicional para párrafos con marcadores específicos
        for paragraph in doc.paragraphs:
            # Buscar si hay algún marcador que no se haya reemplazado
            for key in ['{{GRADO_VO_BO}}', '{{NOMBRE_VO_BO}}', '{{MATRICULA_VO_BO}}']:
                if key in paragraph.text:
                    # Reemplazar directamente
                    valor = ''
                    if key == '{{GRADO_VO_BO}}':
                        valor = data['Grado_Vo_Bo']
                    elif key == '{{NOMBRE_VO_BO}}':
                        valor = data['Nombre_Vo_Bo']
                    elif key == '{{MATRICULA_VO_BO}}':
                        valor = data['Matricula_Vo_Bo']
                    
                    paragraph.text = paragraph.text.replace(key, valor)
                    aplicar_formato_texto(paragraph)

        # Guardar el documento
        output_path = os.path.join(output_dir, template_name + ".docx")
        doc.save(output_path)

        return output_path

    except Exception as e:
        raise Exception(f"Error al crear documento: {str(e)}")