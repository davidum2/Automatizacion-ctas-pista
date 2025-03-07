import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_of_remision(output_dir, data):
    """
    Crea un oficio de remisión basado en los datos de factura.
    
    Args:
        output_dir (str): Directorio donde se guardará el oficio
        data (dict): Datos de la factura
        
    Returns:
        str: Ruta al documento generado
    """
    try:
        # Crear un nuevo documento
        doc = Document()
        
        # Configurar márgenes del documento
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)
        
        # Agregar encabezado
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run("DEPENDENCIA")
        header_run.bold = True
        header_run.font.size = Pt(14)
        
        # Agregar subencabezado
        subheader = doc.add_paragraph()
        subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subheader_run = subheader.add_run("OFICINA")
        subheader_run.bold = True
        subheader_run.font.size = Pt(12)
        
        # Espacio
        doc.add_paragraph()
        
        # Agregar número de oficio
        oficio = doc.add_paragraph()
        oficio.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        oficio_run = oficio.add_run(f"OFICIO No.: {data['No_of_remision']}")
        oficio_run.bold = True
        oficio_run.font.size = Pt(11)
        
        # Agregar fecha
        fecha = doc.add_paragraph()
        fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fecha_run = fecha.add_run(f"México, {data['Fecha_remision']}")
        fecha_run.bold = True
        fecha_run.font.size = Pt(11)
        
        # Agregar asunto
        asunto = doc.add_paragraph()
        asunto.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        asunto_run = asunto.add_run("ASUNTO: ENVÍO DE DOCUMENTACIÓN")
        asunto_run.bold = True
        asunto_run.font.size = Pt(11)
        
        # Espacio
        doc.add_paragraph()
        
        # Destinatario
        destinatario = doc.add_paragraph()
        destinatario_run = destinatario.add_run("C. NOMBRE DEL DESTINATARIO\nCARGO\nP R E S E N T E")
        destinatario_run.bold = True
        destinatario_run.font.size = Pt(11)
        
        # Espacio
        doc.add_paragraph()
        
        # Cuerpo del oficio
        cuerpo = doc.add_paragraph()
        cuerpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cuerpo_texto = f"""Por medio del presente me permito enviar a usted, factura electrónica No. {data['Serie']}{data['Numero']} de fecha {data['Fecha_factura']}, por la cantidad de {data['monto']} ({data['monto']} M.N.), por concepto de {data['Descripcion_partida']}, para que se realice el trámite correspondiente en {data['Empleo_recurso']}, con cargo a la partida {data['No_partida']}, correspondiente al mes de {data['Mes']} del año en curso, asignado con mensaje No. {data['No_mensaje']} de fecha {data['Fecha_mensaje']}."""
        cuerpo_run = cuerpo.add_run(cuerpo_texto)
        cuerpo_run.font.size = Pt(11)
        
        # Espacio
        doc.add_paragraph()
        
        # Saludo final
        saludo = doc.add_paragraph()
        saludo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        saludo_run = saludo.add_run("Sin otro particular, reciba un cordial saludo.")
        saludo_run.font.size = Pt(11)
        
        # Espacio
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Firma
        firma = doc.add_paragraph()
        firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
        firma_run = firma.add_run("ATENTAMENTE\nNOMBRE DEL REMITENTE\nCARGO")
        firma_run.bold = True
        firma_run.font.size = Pt(11)
        
        # Iniciales
        doc.add_paragraph()
        iniciales = doc.add_paragraph()
        iniciales_run = iniciales.add_run("c.c.p. Archivo.\nXXX/xxx")
        iniciales_run.font.size = Pt(8)
        
        # Guardar el documento
        output_path = os.path.join(output_dir, "oficio_de_remision.docx")
        doc.save(output_path)
        
        return output_path
        
    except Exception as e:
        raise Exception(f"Error al crear oficio de remisión: {str(e)}")
