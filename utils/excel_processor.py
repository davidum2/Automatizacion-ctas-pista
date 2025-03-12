import os
from openpyxl import load_workbook
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def procesar_plantillas_de_las_partidas(partida, facturas_info, partida_dir, datos_comunes):
    """
    Procesa plantillas Excel y Word reemplazando valores en celdas específicas
    sin alterar el formato.

    Args:
        partida (dict): Información de la partida
        facturas_info (list): Lista de facturas procesadas
        partida_dir (str): Directorio de la partida
        datos_comunes (dict): Datos comunes para todas las plantillas

    Returns:
        dict: Diccionario con las rutas a los archivos generados
    """
    plantillas = {
        "ingresos": "plantilla_ingresos_egresos.xlsx",
        "facturas": "plantilla_facturas.xlsx",
        "oficio": "plantilla_oficio.docx"
    }

    # Directorio base de plantillas
    templates_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "plantillas")

    # Archivos generados
    archivos_generados = {}

    # Procesar cada plantilla
    try:
        # Procesar plantilla de ingresos/egresos
        ruta_ingresos = procesar_plantilla_ingresos(
            os.path.join(templates_dir, plantillas["ingresos"]),
            partida_dir,
            partida,
            facturas_info,
            datos_comunes
        )
        archivos_generados["ingresos"] = ruta_ingresos

        # Procesar plantilla de facturas
        ruta_facturas = procesar_plantilla_facturas(
            os.path.join(templates_dir, plantillas["facturas"]),
            partida_dir,
            partida,
            facturas_info,
            datos_comunes
        )
        archivos_generados["facturas"] = ruta_facturas

        # Procesar plantilla de oficio (Word)
        ruta_oficio = procesar_plantilla_oficio(
            os.path.join(templates_dir, plantillas["oficio"]),
            partida_dir,
            partida,
            facturas_info,
            datos_comunes
        )
        archivos_generados["oficio"] = ruta_oficio

        return archivos_generados

    except Exception as e:
        raise Exception(f"Error al procesar plantillas de partida: {str(e)}")


def procesar_plantilla_ingresos(template_path, output_dir, partida, facturas_info, datos_comunes):
    """
    Procesa la plantilla de ingresos/egresos.

    Args:
        template_path (str): Ruta a la plantilla Excel
        output_dir (str): Directorio de salida
        partida (dict): Información de la partida
        facturas_info (list): Lista de facturas procesadas
        datos_comunes (dict): Datos comunes

    Returns:
        str: Ruta al archivo generado
    """
    try:
        # Verificar que la plantilla existe
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"No se encontró la plantilla: {template_path}")

        # Cargar la plantilla existente
        wb = load_workbook(template_path)
        ws = wb.active

        # Construir el texto dinámico para A1
        mes = datos_comunes.get('mes_asignado', '').capitalize()
        partida_num = partida.get('numero', '')
        descripcion = partida.get('descripcion', '')

        texto_encabezado = f'Relación de ingresos y egresos correspondientes al mes de {mes} del 2025, de los recursos asignados a la Partida Presupuestal {partida_num} "{descripcion}".'

        # Aplicar valores a celdas específicas
        data = {
            'A1': texto_encabezado,
            'C3': datetime.now().strftime('%d/%m/%Y'),
            'C4': partida_num,
            'C5': descripcion,
            'C6': mes.capitalize(),
            'C7': "2025"  # Año actual o del ejercicio
        }

        # Calcular sumas y montos
        monto_total = sum(float(factura.get('monto', '0').replace('$', '').replace(',', ''))
                          for factura in facturas_info if isinstance(factura, dict))

        # Agregar totales
        data['G15'] = monto_total  # Celda de total (ajustar según la plantilla)

        # Aplicar valores a celdas
        for coordenada, valor in data.items():
            if coordenada in ws:
                ws[coordenada].value = valor

        # Guardar el archivo
        output_path = os.path.join(output_dir, f"Ingresos_Egresos_Partida_{partida_num}.xlsx")
        wb.save(output_path)

        return output_path

    except Exception as e:
        raise Exception(f"Error al procesar plantilla de ingresos: {str(e)}")


def procesar_plantilla_facturas(template_path, output_dir, partida, facturas_info, datos_comunes):
    """
    Procesa la plantilla de facturas.

    Args:
        template_path (str): Ruta a la plantilla Excel
        output_dir (str): Directorio de salida
        partida (dict): Información de la partida
        facturas_info (list): Lista de facturas procesadas
        datos_comunes (dict): Datos comunes

    Returns:
        str: Ruta al archivo generado
    """
    try:
        # Verificar que la plantilla existe
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"No se encontró la plantilla: {template_path}")

        # Cargar la plantilla existente
        wb = load_workbook(template_path)
        ws = wb.active

        # Construir el texto dinámico para A1
        mes = datos_comunes.get('mes_asignado', '').capitalize()
        partida_num = partida.get('numero', '')
        descripcion = partida.get('descripcion', '')

        texto_encabezado = f'Relación de facturas correspondientes al mes de {mes} del 2025, de los recursos asignados a la Partida Presupuestal {partida_num} "{descripcion}".'

        # Aplicar valores a celdas específicas
        data = {
            'A1': texto_encabezado,
            'B3': partida_num,
            'B4': descripcion,
            'B5': mes.capitalize(),
            'B6': "2025"  # Año actual o del ejercicio
        }

        # Aplicar valores a celdas
        for coordenada, valor in data.items():
            if coordenada in ws:
                ws[coordenada].value = valor

        # Agregar información de facturas
        fila_inicial = 9  # Fila donde comienza la tabla de facturas (ajustar según la plantilla)
        for i, factura in enumerate(facturas_info):
            if not isinstance(factura, dict):
                continue

            fila = fila_inicial + i
            # Extraer fecha de factura y convertirla si es necesario
            fecha_factura = factura.get('fecha', '')
            if isinstance(fecha_factura, str) and '-' in fecha_factura:
                try:
                    fecha_obj = datetime.strptime(fecha_factura, '%Y-%m-%d')
                    fecha_factura = fecha_obj.strftime('%d/%m/%Y')
                except:
                    pass  # Mantener el formato original si hay error

            # Aplicar datos de la factura a las celdas
            ws[f'A{fila}'] = i + 1  # Número secuencial
            ws[f'B{fila}'] = fecha_factura
            ws[f'C{fila}'] = factura.get('emisor', '')
            ws[f'D{fila}'] = factura.get('conceptos', '')
            ws[f'E{fila}'] = factura.get('serie_numero', '')
            ws[f'F{fila}'] = factura.get('rfc_emisor', '')

            # Convertir monto a número si viene como string
            monto_str = factura.get('monto', '0')
            if isinstance(monto_str, str):
                monto_str = monto_str.replace('$', '').replace(',', '')
                try:
                    monto = float(monto_str)
                except:
                    monto = 0
            else:
                monto = float(monto_str)

            ws[f'G{fila}'] = monto

        # Calcular el total de montos
        fila_total = fila_inicial + len(facturas_info) + 1
        formula_total = f"=SUM(G{fila_inicial}:G{fila_inicial + len(facturas_info) - 1})"
        ws[f'G{fila_total}'] = formula_total

        # Guardar el archivo
        output_path = os.path.join(output_dir, f"Relacion_Facturas_Partida_{partida_num}.xlsx")
        wb.save(output_path)

        return output_path

    except Exception as e:
        raise Exception(f"Error al procesar plantilla de facturas: {str(e)}")


def procesar_plantilla_oficio(template_path, output_dir, partida, facturas_info, datos_comunes):
    """
    Procesa la plantilla de oficio en formato Word.

    Args:
        template_path (str): Ruta a la plantilla Word
        output_dir (str): Directorio de salida
        partida (dict): Información de la partida
        facturas_info (list): Lista de facturas procesadas
        datos_comunes (dict): Datos comunes

    Returns:
        str: Ruta al archivo generado
    """
    try:
        # Verificar que la plantilla existe
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"No se encontró la plantilla: {template_path}")

        # Cargar la plantilla
        doc = Document(template_path)

        # Datos para reemplazar
        mes = datos_comunes.get('mes_asignado', '').capitalize()
        partida_num = partida.get('numero', '')
        descripcion = partida.get('descripcion', '')
        fecha_doc = datos_comunes.get('fecha_documento_texto', '')

        # Total de facturas y monto total
        total_facturas = len(facturas_info)
        monto_total = sum(float(factura.get('monto', '0').replace('$', '').replace(',', ''))
                          for factura in facturas_info if isinstance(factura, dict))
        monto_formateado = "$ {:,.2f}".format(monto_total)

        # Buscar marcadores en el documento y reemplazarlos con los datos
        for paragraph in doc.paragraphs:
            if '{{FECHA_DOCUMENTO}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{FECHA_DOCUMENTO}}', fecha_doc)
            if '{{MES}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{MES}}', mes)
            if '{{PARTIDA}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{PARTIDA}}', partida_num)
            if '{{DESCRIPCION}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{DESCRIPCION}}', descripcion)
            if '{{TOTAL_FACTURAS}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{TOTAL_FACTURAS}}', str(total_facturas))
            if '{{MONTO_TOTAL}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{MONTO_TOTAL}}', monto_formateado)

        # También buscar en las tablas si existen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{{FECHA_DOCUMENTO}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{FECHA_DOCUMENTO}}', fecha_doc)
                        if '{{MES}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{MES}}', mes)
                        if '{{PARTIDA}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{PARTIDA}}', partida_num)
                        if '{{DESCRIPCION}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{DESCRIPCION}}', descripcion)
                        if '{{TOTAL_FACTURAS}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{TOTAL_FACTURAS}}', str(total_facturas))
                        if '{{MONTO_TOTAL}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{{MONTO_TOTAL}}', monto_formateado)

        # Guardar el documento
        output_path = os.path.join(output_dir, f"Oficio_Resumen_Partida_{partida_num}.docx")
        doc.save(output_path)

        return output_path

    except Exception as e:
        raise Exception(f"Error al procesar plantilla de oficio: {str(e)}")
